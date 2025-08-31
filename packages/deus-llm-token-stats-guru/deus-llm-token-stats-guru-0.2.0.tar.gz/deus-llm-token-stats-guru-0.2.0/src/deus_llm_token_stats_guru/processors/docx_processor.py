"""DOCX file processor."""

from pathlib import Path
from typing import Set, List
import logging

from .base import BaseFileProcessor
from .._exceptions import FileProcessingError

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseFileProcessor):
    """Processor for DOCX files."""
    
    @property
    def supported_extensions(self) -> Set[str]:
        """Return supported file extensions."""
        return {".docx", ".doc"}
    
    @property
    def processor_name(self) -> str:
        """Return processor name."""
        return "DOCX"
    
    def extract_text(self, file_path: Path) -> List[str]:
        """Extract text content from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of text strings (paragraphs)
            
        Raises:
            FileProcessingError: If DOCX cannot be processed
        """
        try:
            # Try python-docx first
            try:
                from docx import Document
                
                doc = Document(file_path)
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_content.append(text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                return text_content
                
            except ImportError:
                # Fallback for .doc files using python-docx2txt
                try:
                    import docx2txt
                    
                    text = docx2txt.process(file_path)
                    if text and text.strip():
                        # Split by paragraphs (double newlines)
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        return paragraphs if paragraphs else [text]
                    else:
                        return []
                        
                except ImportError:
                    raise FileProcessingError(
                        "No DOCX processing library available. "
                        "Please install python-docx or docx2txt"
                    )
        
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from DOCX: {e}") from e