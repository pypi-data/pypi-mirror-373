from docx import Document
import re

def bold_words_in_text(input_docx_path, bold_words, output_path='output.docx'):
    document = Document(input_docx_path)
    # Normalize phrases (lowercase, strip spaces)
    phrases_to_bold = [phrase.strip() for phrase in bold_words]

    for para in document.paragraphs:
        # Rebuild full text and mapping of character positions to runs
        full_text = ""
        char_to_run = []
        for run in para.runs:
            for i, char in enumerate(run.text):
                full_text += char
                char_to_run.append((run, i))

        # Unbold all runs first
        for run in para.runs:
            run.bold = False

        lowered_text = full_text.lower()
        bold_flags = [False] * len(full_text)

        # Mark characters to be bolded using strict matching
        for phrase in phrases_to_bold:
            escaped_phrase = re.escape(phrase)
            # Only block letters before/after, allow numbers, apostrophes, and hyphens
            pattern = rf"(?i)(?<![A-Za-z]){escaped_phrase}(?![A-Za-z])"
            
            for match in re.finditer(pattern, lowered_text):
                start_idx, end_idx = match.span()
                for i in range(start_idx, end_idx):
                    bold_flags[i] = True

        # Apply bolding by splitting runs where needed
        new_runs = []
        i = 0
        while i < len(char_to_run):
            run, char_idx = char_to_run[i]
            start_i = i
            is_bold = bold_flags[i]

            while (i < len(char_to_run)
                   and char_to_run[i][0] == run
                   and bold_flags[i] == is_bold):
                i += 1

            text_chunk = run.text[char_to_run[start_i][1]:char_to_run[i-1][1]+1]
            new_runs.append((text_chunk, is_bold))

        # Clear the paragraph and rebuild runs
        para.clear()
        for text, bold_flag in new_runs:
            if text:
                new_run = para.add_run(text)
                new_run.bold = bold_flag

    # Save the new document
    document.save(output_path)
    print(f"Document has been saved to this path: '{output_path}'")
    return


if __name__ == "__main__":
    input_path = "/home/cait/Desktop/Python-Package/DocxFiles/Python Docs Text.docx"
    bold_words = [
    "prime minister",
    "unity",
    "economic growth",
    "leadership",
    "real change",
    "it's",
    "science"
    ]
    bold_words_in_text(input_path, bold_words)