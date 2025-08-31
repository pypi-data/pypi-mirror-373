import typer, re, json, csv, ipaddress, os, tldextract
from pathlib import Path
from typing import List, Dict 
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from docx import Document

app = typer.Typer()


IOC_PATTERNS = {
    "ipv4": re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b"),
    "md5": re.compile(r"(?<![a-f0-9])[a-f0-9]{32}(?![a-f0-9])", re.IGNORECASE),
    "sha1": re.compile(r"(?<![a-f0-9])[a-f0-9]{40}(?![a-f0-9])", re.IGNORECASE),
    "sha224": re.compile(r"(?<![a-f0-9])[a-f0-9]{56}(?![a-f0-9])", re.IGNORECASE),
    "sha256": re.compile(r"(?<![a-f0-9])[a-f0-9]{64}(?![a-f0-9])", re.IGNORECASE),
     "sha384": re.compile(r"(?<![a-f0-9])[a-f0-9]{96}(?![a-f0-9])", re.IGNORECASE),
    "sha512": re.compile(r"(?<![a-f0-9])[a-f0-9]{128}(?![a-f0-9])", re.IGNORECASE),
    "domain_standard": re.compile(r"\b(?:[a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}\b"),
    "domain_suspicious": re.compile(r"\b(?:[a-zA-Z0-9\-_]+\.)+(?:exe|dll|scr|pdf|docx?|xlsx?|js|vbs|bat|cmd|ps1|lnk|zip|rar|7z)\.[a-zA-Z]{2,}\b", re.IGNORECASE),
    "url": re.compile(r"https?://(?:(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+[A-Z]{2,})|(?:(?:\d{1,3}\.){3}\d{1,3}))(?::\d+)?(?:/[\w\-\.~!$&'()*+,;=:@%]*)?(?=\s|$)", re.IGNORECASE),
    "email": re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"),
    "cve": re.compile(r"\bCVE-\d{4}-\d{4,7}\b", re.IGNORECASE)
}

CATEGORY_MAPPING = {
    "ipv4": "IP Addresses",
    "md5": "Hashes",
    "sha1": "Hashes",
    "sha224": "Hashes",
    "sha256": "Hashes",
    "sha384": "Hashes",
    "sha512": "Hashes",
    "domain_standard": "Domains",
    "domain_suspicious": "Suspicious Domains",
    "url": "URLs",
    "email": "Emails",
    "cve": "CVEs"
}

def extract_iocs_from_text(text: str, patterns: Dict[str, re.Pattern]) -> Dict[str, List[str]]:
    results = {}
    for ioc_type, pattern in patterns.items():
        found_items = pattern.findall(text)
        results[ioc_type] = sorted(list(set(found_items)))
    return results

def read_text_content(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return file_path.read_text(encoding="cp1252")
        except UnicodeDecodeError:
            return file_path.read_text(encoding="utf-8", errors='ignore')
        
def read_large_text_content(file_path: Path) -> str:
    text_lines = []
    encodings = ['utf-8', 'cp1252', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                for line in f:
                    text_lines.append(line)
            return "".join(text_lines)
        except UnicodeDecodeError:
            text_lines = []
            continue
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                text_lines.append(line)
        return "".join(text_lines)
    except Exception as e:
        raise Exception(f"Failed to read text file: {e}")
    
def extract_iocs_from_text_streaming(file_path: Path, patterns: Dict[str, re.Pattern]) -> Dict[str, List[str]]:
    results = {ioc_type: set() for ioc_type in patterns.keys()}
    encodings = ['utf-8', 'cp1252', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                for line in f:
                    for ioc_type, pattern in patterns.items():
                        found_items = pattern.findall(line)
                        results[ioc_type].update(found_items)
            break
        except UnicodeDecodeError:
            results = {ioc_type: set() for ioc_type in patterns.keys()}
            continue
    else:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    for ioc_type, pattern in patterns.items():
                        found_items = pattern.findall(line)
                        results[ioc_type].update(found_items)
        except Exception as e:
            raise Exception(f"Failed to read text file: {e}")
    
    return {ioc_type: sorted(list(items)) for ioc_type, items in results.items()}

def read_json_content(file_path: Path) -> str:
    try:
        import json
        data = json.loads(file_path.read_text(encoding="utf-8"))
        def extract_values(obj):
            if isinstance(obj, dict):
                for value in obj.values():
                    yield from extract_values(value)
            elif isinstance(obj, list):
                for item in obj:
                    yield from extract_values(item)
            else:
                yield str(obj)
        return " ".join(extract_values(data))
    except Exception as e:
        raise Exception(f"Failed to read JSON: {e}")
        
def read_csv_content(file_path: Path) -> str:
    try:
        text = ""
        encodings = ['utf-8', 'cp1252', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        text += " ".join(row) + "\n"
                return text
            except UnicodeDecodeError:
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore', newline='') as f:
            reader = csv.reader(f)
            for row in reader:
                text += " ".join(row) + "\n"
        return text
    except Exception as e:
        raise Exception(f"Failed to read CSV: {e}")

def read_pdf_content(file_path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
        text = ""
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Failed to read PDF: {e}")

def read_html_content(file_path: Path) -> str:
    try:
        html_content = file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html_content, 'lxml')
        text = soup.get_text()
        return text
    except UnicodeDecodeError:
        try:
            html_content = file_path.read_text(encoding="cp1252")
            soup = BeautifulSoup(html_content, 'lxml')
            text = soup.get_text()
            return text
        except UnicodeDecodeError:
            html_content = file_path.read_text(encoding="utf-8", errors='ignore')
            soup = BeautifulSoup(html_content, 'lxml')
            text = soup.get_text()
            return text
    except Exception as e:
        raise Exception(f"Failed to read HTML: {e}")
    
def read_docx_content(file_path: Path) -> str:
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Failed to read DOCX: {e}")

def read_file_content(file_path: Path) -> str:
    extension = file_path.suffix.lower()[1:]
    if extension == 'pdf':
        return read_pdf_content(file_path)
    elif extension == 'html': 
        return read_html_content(file_path)
    elif extension == 'docx':
        return read_docx_content(file_path)
    elif extension == 'json':
        return read_json_content(file_path)
    elif extension == 'csv':
        return read_csv_content(file_path)
    else:
        return read_text_content(file_path)
    
def validate_ipv4(ip: str) -> bool:
    try:
        ipaddress.IPv4Address(ip)
        return True
    except ipaddress.AddressValueError:
        return False
    
def validate_domain(domain: str) -> bool:
    if validate_ipv4(domain):
        return False
        
    extracted = tldextract.extract(domain)
    return bool(extracted.suffix and extracted.domain)
    
def validate_iocs(results: Dict[str, List[str]]) -> Dict[str, List[str]]:
    validated_results = {}
    
    for ioc_type, items in results.items():
        if ioc_type == "ipv4":
            validated_results[ioc_type] = [ip for ip in items if validate_ipv4(ip)]
        elif ioc_type in ["domain_standard", "domain_suspicious"]:
            validated_results[ioc_type] = [domain for domain in items if validate_domain(domain)]
        else:
            validated_results[ioc_type] = items
    
    return validated_results

SUPPORTED_FORMATS = {"txt", "pdf", "html", "log", "csv", "json", "docx"}

@app.command(help="Extract IOCs from a file")
def extract(
    file_path: str = typer.Argument(..., help="Path to the input file"),
    output: str = typer.Option(None, "--output", "-o", help="Path to the output file (Optional)"),
    format: str = typer.Option("text", "--format", "-f", help="Output format: text, json or csv"),
    streaming: bool = typer.Option(False, "--streaming", "-s", help="Use streaming mode for large files"),
    ipv4: bool = typer.Option(False, "--ipv4", help="Extract only IPv4 addresses"),
    md5: bool = typer.Option(False, "--md5", help="Extract only MD5 hashes"),
    sha1: bool = typer.Option(False, "--sha1", help="Extract only SHA-1 hashes"),
    sha256: bool = typer.Option(False, "--sha256", help="Extract only SHA-256 hashes"),
    domain: bool = typer.Option(False, "--domain", help="Extract only domains"),
    url: bool = typer.Option(False, "--url", help="Extract only URLs"),
    email: bool = typer.Option(False, "--email", help="Extract only emails"),
    cve: bool = typer.Option(False, "--cve", help="Extract only CVEs"),
    hash: bool = typer.Option(False, "--hash", help="Extract all hash types (MD5, SHA-1, SHA-256, etc.)")
    ):

    file_path_obj = Path(file_path)

    if not file_path_obj.exists():
        typer.echo(f"Error: File '{file_path}' does not exist.")
        raise typer.Exit(code=1)
    
    if not file_path_obj.is_file():
        typer.echo(f"Error: '{file_path}' is a directory, not a file.")
        raise typer.Exit(code=1)
    
    suffix = file_path_obj.suffix.lower()
    if not suffix:
        typer.echo(f"Error: File '{file_path}' has no extension.")
        raise typer.Exit(code=1)
    
    if file_path_obj.suffix.lower()[1:] not in SUPPORTED_FORMATS:
        typer.echo("Error: Unsupported file format. Use 'formats' command to see supported formats.")
        raise typer.Exit(code=1)
    
    if not os.access(file_path_obj, os.R_OK):
        typer.echo(f"Error: Cannot read file '{file_path}'. Check permissions.")
        raise typer.Exit(code=1)
    
    try:
        if streaming and file_path_obj.suffix.lower()[1:] in {'txt', 'log'}:
            results = extract_iocs_from_text_streaming(file_path_obj, IOC_PATTERNS)
        else:
            content = read_file_content(file_path_obj)
            results = extract_iocs_from_text(content, IOC_PATTERNS)
    except Exception as e:
        typer.echo(f"Error reading file: {e}")
        raise typer.Exit(code=1)
    
    if any([ipv4, md5, sha1, sha256, domain, url, email, cve, hash]):
        filtered_results = {}
        
        included_types = set()
        
        if ipv4:
            included_types.add("ipv4")
        if md5:
            included_types.add("md5")
        if sha1:
            included_types.add("sha1")
        if sha256:
            included_types.add("sha256")
        if domain:
            included_types.update(["domain_standard", "domain_suspicious"])
        if url:
            included_types.add("url")
        if email:
            included_types.add("email")
        if cve:
            included_types.add("cve")
        if hash:
            included_types.update(["md5", "sha1", "sha224", "sha256", "sha384", "sha512"])
        
        for ioc_type, items in results.items():
            if ioc_type in included_types:
                filtered_results[ioc_type] = items
        
        results = filtered_results    

    results = validate_iocs(results)



    total_found = sum(len(items) for items in results.values())
    if format == "json":
        output_content = json.dumps(results, indent=2)
    elif format == "csv":
        output_content = "CATEGORY,SUBCATEGORY,VALUE\n"
    
        sorted_categories = sorted(results.keys())
    
        for ioc_type in sorted_categories:
            items = results[ioc_type]
            if items:
                main_category = CATEGORY_MAPPING.get(ioc_type, "Other")
                subcategory = ioc_type
            
                if ioc_type in ["md5", "sha1", "sha224", "sha256", "sha384", "sha512"]:
                    subcategory = ioc_type.upper()
            
                for item in items:
                    output_content += f"{main_category},{subcategory},{item}\n"
    else:
        output_content = ""
        for ioc_type, items in results.items():
            if items:
                category_name = CATEGORY_MAPPING.get(ioc_type, ioc_type.upper())
                output_content += f"\n{category_name}:\n"
                for item in items:
                    output_content += f"  {item}\n"
    
    output_content += f"\nTotal IOCs found: {total_found}\n"

    if output:
        try:
            with open(output, 'w', encoding='utf-8') as f:
                f.write(output_content)
            typer.echo(f"Results saved to: {output}")
        except Exception as e:
            typer.echo(f"Error writing to file: {e}")
            raise typer.Exit(code=1)
    else:
        typer.echo(output_content)



@app.command(help="List of available formats to extract IOCs")
def formats():
    typer.echo("Supported formats for IOC extraction:")
    for format_name in SUPPORTED_FORMATS:
        typer.echo(f"- {format_name}")

@app.command(help="List all IOC types that can be extracted")
def types():
    typer.echo("Supported IOC types:")
    for ioc_type in IOC_PATTERNS.keys():
        typer.echo(f"- {ioc_type}")



app()